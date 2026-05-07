"""
텀프로젝트 계획서 Word 문서 생성
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_PATH = "/Users/jin/석사논문/TAVI/응용지리정보체계론_텀프로젝트/docs/2026-05-11_텀프로젝트계획서.docx"

doc = Document()

# ── 페이지 여백 설정 (A4, 상하좌우 2.5cm) ────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 기본 폰트 설정 함수 ──────────────────────────────────────────────────────
def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size = 13 if level == 1 else 11
    set_font(run, size=size, bold=True, color=(31, 73, 125) if level == 1 else (0, 0, 0))
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.7)
    if bold_part and bold_part in text:
        before, after = text.split(bold_part, 1)
        if before:
            r = p.add_run(before); set_font(r)
        r = p.add_run(bold_part); set_font(r, bold=True)
        if after:
            r = p.add_run(after); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p

# ════════════════════════════════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("거리 영상 기반 MLLM MRT 추정을 활용한\n대중교통 Thermal Catchment 방법론 제안")
set_font(run, size=15, bold=True, color=(31, 73, 125))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(10)
run2 = p2.add_run("응용지리정보체계론 Final Project 계획서  |  신진  |  2026-05-11")
set_font(run2, size=10, color=(89, 89, 89))

doc.add_paragraph()  # 구분선 대용 공백

# ════════════════════════════════════════════════════════════════════════════
# 1. 필요성 및 연구배경
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. 필요성 및 연구배경")

add_body(doc,
    "기후변화에 따른 도시 폭염 빈도·강도의 증가로 보행 열 환경이 대중교통 이용 행태에 "
    "미치는 영향이 주목받고 있다. 역세권 보행 접근성은 단순한 거리·시간이 아닌 열 노출 "
    "수준에 따라 실질적으로 제약받는다. 이를 정량화하기 위해서는 링크(도로 구간)별 "
    "평균복사온도(MRT) 또는 UTCI 산출이 필요하다."
)

add_body(doc, "그러나 기존의 물리 기반 MRT/UTCI 산출 과정에는 다음과 같은 어려움이 따른다.")

# 표 — 문제점
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
headers = ["문제", "내용"]
rows_data = [
    ("데이터 취득",    "DSM(건물+수목 높이), SVF, 기상 센서망 등 다중 소스 필요"),
    ("방법론 복잡성",  "직산 분리 → 장파 추정 → Stefan-Boltzmann 역산 → UTCI 계산"),
    ("간략화 오류",   "LiDAR 미확보 시 합성 DSM 사용, 기상 공간 단일값 적용 등"),
    ("확장 한계",     "성동구 수준에서도 구축 비용 과다; 전국 확장 시 현실적으로 불가"),
]
for i, (h, _) in enumerate([(headers[0], ""), (headers[1], "")]):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(headers[i])
    set_font(run, bold=True, size=10)
for i, (k, v) in enumerate(rows_data):
    r = table.rows[i+1]
    run_k = r.cells[0].paragraphs[0].add_run(k); set_font(run_k, size=10)
    run_v = r.cells[1].paragraphs[0].add_run(v); set_font(run_v, size=10)

doc.add_paragraph()

add_body(doc,
    "반면, MLLM(Multimodal Large Language Model)은 거리 영상(Google Street View)으로부터 "
    "건물 형태, 수목 캐노피, 차양 구조물, 노면 재질 등 MRT 결정 요소를 직접 추론할 수 있다. "
    "본 연구는 MLLM을 활용하여 복잡한 물리 기반 MRT 산출 과정(DSM→SVF→SOLWEIG)을 거리 영상 "
    "추정으로 대체하는 방법론을 제안하고, 이를 대중교통(지하철역·버스 정류장) Thermal Catchment "
    "프레임워크에 적용하는 것을 목적으로 한다."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. 관련 연구 리뷰
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. 관련 연구 리뷰")

add_heading(doc, "■ 도시 열 환경 및 보행 접근성", level=2)
add_bullet(doc, "Lindberg & Grimmond (2011): SVF 기반 SOLWEIG 모델로 링크별 MRT 산출 — 물리적으로 엄밀하나 고해상도 DSM 의존")
add_bullet(doc, "Bröde et al. (2012): UTCI 공식 등급 체계 제시; ≥38°C를 'very strong heat stress'로 정의")
add_bullet(doc, "Thermal Catchment 개념(저자 선행 연구): UTCI 기반 Hard Cut을 적용한 역세권 도보 도달권 축소 분석")

add_heading(doc, "■ 거리 영상 + MLLM 기반 도시 속성 추출", level=2)
add_bullet(doc, "Li et al. (StreetViewLLM): GSV + Chain-of-Thought MLLM으로 5가지 도시 지표 추출; 7개 도시 검증, 기준 모델 대비 49% 이상 성능 향상 — CoT 프롬프팅으로 복잡한 도시 속성 추론 가능함을 직접 입증", bold_part="CoT 프롬프팅으로 복잡한 도시 속성 추론 가능함을 직접 입증")
add_bullet(doc, "Liang et al. (OpenFACADES, 2025): GSV + VLM으로 건물 외관 속성(재질·H/W비 등) 대규모 자동 추출; 7개 도시 120만 건물 데이터셋 구축 — 본 연구 프롬프트 항목의 직접적 방법론 근거", bold_part="본 연구 프롬프트 항목의 직접적 방법론 근거")

add_heading(doc, "■ MLLM 기반 도시 장면 주관적 평가", level=2)
add_bullet(doc, "He et al. (UrbanFeel, 2025): 14.3K VQA 벤치마크; 안전성·쾌적성 등 주관적 환경 인식에서 20개 MLLM 평가 — MLLM이 보행환경 쾌적도를 시각적으로 추론할 수 있음을 실험 검증")
add_bullet(doc, "Feng et al. (UrbanLLaVA, 2025): 위성·거리뷰·지리공간·궤적 통합 Urban MLLM — 도시 거리 영상 기반 다중 태스크 추론 가능성 제시")
add_bullet(doc, "Zhang et al. (EarthGPT, IEEE TGRS 2024): 원격탐사 전용 MLLM — 도메인 특화 학습의 성능 우위 입증; 향후 열환경 평가 전용 모델 fine-tuning 가능성 시사 (저자 발제 논문)")

# ════════════════════════════════════════════════════════════════════════════
# 3. 연구지역 및 연구방법
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. 연구지역 및 연구방법")

add_heading(doc, "■ 연구지역", level=2)
add_body(doc,
    "서울특별시 성동구 — 지하철역 7개소(응봉·성수·왕십리·행당·뚝섬·서울숲·옥수) + 버스 정류장. "
    "기존 TAVI 석사 연구와 동일 지역으로, 선행 분석 결과와의 직접 연계가 가능하다."
)

add_heading(doc, "■ 연구방법", level=2)

steps = [
    ("Step 1. 거리 영상 수집",
     "Google Street View Static API를 활용하여 역세권 보행 네트워크 링크 중심점 좌표 기반 크롤링. "
     "4방향(0°·90°·180°·270°) 이미지, 8월 촬영분 필터링(Metadata API 활용)."),
    ("Step 2. MLLM 기반 링크별 MRT 추정",
     "Qwen2-VL-7B(1차) 또는 GPT-4o(비교)를 활용. 구조화 프롬프트로 그늘 비율·수목 캐노피·"
     "건물 H/W비·차양 구조물·노면 재질을 항목별 평가하여 MRT를 추정. "
     "물리 기반 DSM→SVF→SOLWEIG 산출 과정을 대체."),
    ("Step 3. 대중교통 Thermal Catchment 적용",
     "MLLM MRT 추정값을 링크별 통행 가중치 또는 Hard Cut 기준으로 변환. "
     "기존 TAVI 프레임워크(networkx 기반) 재사용하여 지하철역·버스 정류장 대상 "
     "Classic vs. Thermal Catchment 비교."),
]
for title, body in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(title + "\n"); set_font(r, bold=True, size=10.5)
    r2 = p.add_run(body);        set_font(r2, size=10.5)

# ════════════════════════════════════════════════════════════════════════════
# 4. 활용 예정 MLLM 모델 및 프롬프트 초안
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. 활용 예정 MLLM 모델 및 프롬프트 초안")

add_heading(doc, "■ 모델 비교", level=2)

table2 = doc.add_table(rows=5, cols=4)
table2.style = "Table Grid"
h2 = ["모델", "제공사", "특징", "비용"]
for i, h in enumerate(h2):
    r = table2.rows[0].cells[i].paragraphs[0].add_run(h)
    set_font(r, bold=True, size=10)
model_rows = [
    ("GPT-4o",           "OpenAI",    "최고 수준 이미지 이해·수치 추론",         "API 유료"),
    ("Claude 3.5 Sonnet","Anthropic", "구조화 출력 우수, 한국어 강점",            "API 유료"),
    ("Qwen2-VL-7B",      "Alibaba",   "오픈소스, 수업 실습 경험",                "무료(Colab)"),
    ("Gemini 1.5 Pro",   "Google",    "스트리트뷰 연계 가능성",                  "API 유료"),
]
for i, row_data in enumerate(model_rows):
    for j, val in enumerate(row_data):
        run = table2.rows[i+1].cells[j].paragraphs[0].add_run(val)
        bold = (i == 2)  # Qwen2-VL 강조
        set_font(run, size=10, bold=bold)

doc.add_paragraph()
add_body(doc, "→ 1차 실험: Qwen2-VL-7B (수업 실습 경험, Colab 무료 실행 가능) / 비교 실험: GPT-4o")

add_heading(doc, "■ 프롬프트 초안 (열 노출 평가용)", level=2)

prompt_lines = [
    "당신은 도시 열 환경 전문 평가원입니다.",
    "아래 거리 사진은 [서울, 여름 8월 오후 1시 기준 폭염일 조건]의 보행 경로입니다.",
    "사진에서 직접 보이는 것만 근거로 각 항목을 평가하세요.",
    "",
    "【그늘 비율】: 현재 경로에서 그늘이 드리운 면적 비율 추정 (0~100%)",
    "【수목 캐노피】: 가로수·녹지에 의한 차양 정도 (없음/낮음/중간/높음)",
    "【건물 차폐】: 건물에 의한 태양 차폐 정도, H/W비 추정 (개활지/반차폐/완전협곡)",
    "【노면 재질】: 아스팔트/콘크리트/블록 등 (복사열 흡수에 영향)",
    "【차양 구조물】: 아케이드, 캐노피, 파라솔 등 인공 차양 유무",
    "【열 노출 점수】: X.X점 (1.0=극고노출 ~ 10.0=완전차폐·쾌적)",
    "【MRT 추정 범주】: 낮음(<45°C) / 중간(45~55°C) / 높음(>55°C)",
    "【종합 평가】: 위 항목을 종합한 열 노출 환경 한 문장",
    "",
    "※ 모든 항목 필수 작성. 확인 불가 시 '확인 불가' 기재.",
    "※ 【열 노출 점수】는 반드시 1.0~10.0 소수점 한 자리 숫자 기입.",
]
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.5)
p.paragraph_format.space_before = Pt(2)
for line in prompt_lines:
    run = p.add_run(line + "\n")
    set_font(run, size=9.5)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

# ════════════════════════════════════════════════════════════════════════════
# 5. 기대 효과 및 의의
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. 기대 효과 및 의의")

effects = [
    ("데이터 구축 비용 절감",  "DSM, SVF, 기상 센서망 없이 무료 거리 영상으로 대체"),
    ("시각적 맥락 통합",       "물리 공식이 놓치는 요소(아케이드, 지하 통로, 차양막) 포착 가능"),
    ("적용 가능성",            "별도 물리 모델 없이 성동구 전 역세권 보행 링크에 바로 적용 가능"),
    ("연구 가성비",            "기존 Thermal Catchment 프레임워크 재사용, MRT 산출 단계만 교체"),
    ("석사 논문 연계",         "성동구 비교 분석으로 TAVI 방법론 고도화 기여"),
]
for i, (k, v) in enumerate(effects):
    add_bullet(doc, f"{k}: {v}", bold_part=k)

# ════════════════════════════════════════════════════════════════════════════
# 6. 참고문헌
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. 참고문헌")

refs = [
    "Bröde, P., et al. (2012). Deriving the operational procedure for the Universal Thermal Climate Index (UTCI). International Journal of Biometeorology, 56(3), 481–494.",
    "Feng, J., Wang, S., Liu, T., Xi, Y., & Li, Y. (2025). UrbanLLaVA: A Multi-modal Large Language Model for Urban Intelligence with Spatial Reasoning and Understanding. arXiv preprint.",
    "He, J., et al. (2025). UrbanFeel: A Comprehensive Benchmark for Temporal and Perceptual Understanding of City Scenes through Human Perspective. arXiv:2509.22228.",
    "Liang, X., Xie, J., Zhao, T., Stouffs, R., & Biljecki, F. (2025). OpenFACADES: An open framework for architectural caption and attribute data enrichment via street view imagery.",
    "Li, Z., Xu, J., Wang, S., Wu, Y., & Li, H. StreetViewLLM: Extracting Geographic Information Using a Chain-of-Thought Multimodal Large Language Model. arXiv preprint.",
    "Lindberg, F., & Grimmond, C. S. B. (2011). The influence of vegetation and building morphology on shadow patterns and mean radiant temperatures in urban areas. Theoretical and Applied Climatology, 105(3–4), 311–323.",
    "Zhang, W., Cai, M., Zhang, T., Zhuang, Y., & Mao, X. (2024). EarthGPT: A Universal Multimodal Large Language Model for Multisensor Image Comprehension in Remote Sensing Domain. IEEE Transactions on Geoscience and Remote Sensing, 62, 5917820.",
    "Höppe, P. (1992). A new procedure to determine the mean radiant temperature outdoors. Meteorologische Zeitschrift, 1(3), 147–148.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before    = Pt(1)
    p.paragraph_format.space_after     = Pt(1)
    p.paragraph_format.left_indent     = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    run = p.add_run(ref)
    set_font(run, size=9.5)

doc.save(OUT_PATH)
print(f"저장 완료: {OUT_PATH}")
